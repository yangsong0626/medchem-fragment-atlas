#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
PROJECT = ROOT.parent
SUMMARY = ROOT / "manuscript" / "manuscript_summary.json"
OUT = ROOT / "manuscript" / "medchem_fragment_atlas_preprint.docx"
FIG = ROOT / "figures"
PNG = FIG / "png"


def load_summary() -> dict:
    return json.loads(SUMMARY.read_text(encoding="utf-8"))


def fmt(value, digits: int = 2) -> str:
    if value is None:
        return "NA"
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if abs(number) >= 100:
        return f"{number:.0f}"
    if abs(number) >= 10:
        return f"{number:.1f}"
    return f"{number:.{digits}f}"


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.5) -> None:
    path = PNG / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
    add_paragraph(doc, caption)


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(18)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"


def build() -> None:
    data = load_summary()
    c = data["counts"]
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("A BRICS Fragment Atlas of Parent-Molecule ADMET Associations from ChEMBL and TDC")

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("Song Yang and MedChem Fragment Atlas Contributors").bold = True
    affiliation = doc.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.add_run("Preprint draft prepared for arXiv submission")

    doc.add_heading("Abstract", level=1)
    add_paragraph(
        doc,
        "Medicinal chemists often reason about structural fragments when optimizing absorption, distribution, metabolism, excretion, and toxicity "
        "(ADMET), yet public ADMET resources are usually organized around whole molecules or assay records rather than reusable fragment motifs. "
        "We present MedChem Fragment Atlas, a reproducible BRICS-based pipeline and interactive web application that decomposes ChEMBL molecules into "
        "canonical attachment-aware fragments, maps parent-molecule physicochemical and ADMET measurements onto those fragments, and summarizes clean "
        "Therapeutics Data Commons (TDC) ADMET benchmarks as fragment-associated endpoint distributions. In the current prototype, the atlas contains "
        f"{c['molecules']:,} ChEMBL parent molecules, {c['fragments']:,} ChEMBL BRICS fragments, {c['tdc_admet_measurements']:,} clean TDC ADMET "
        f"measurements, {c['tdc_fragments']:,} TDC BRICS fragments, and {c['tdc_fragment_admet_stats']:,} clean fragment-endpoint rollups. "
        "To address common reviewer concerns about uncertainty and confounding, we add bootstrap confidence intervals for fragment-endpoint medians, "
        "nearest-neighbor matched-context validation within individual TDC tasks, and automatically generated medicinal chemistry case studies. "
        "The resource is intended as a fragment-centric hypothesis generator: reported values are aggregates of parent molecules containing a fragment, "
        "not intrinsic or measured properties of the isolated fragment."
    )

    doc.add_heading("Keywords", level=1)
    add_paragraph(doc, "BRICS decomposition; ChEMBL; Therapeutics Data Commons; ADMET; fragment analysis; matched context; medicinal chemistry informatics.")

    doc.add_heading("1. Introduction", level=1)
    add_paragraph(
        doc,
        "Fragment-level reasoning is a central part of medicinal chemistry. Chemists routinely ask whether a piperazine, phenyl linker, carboxylic acid, "
        "trifluoromethyl group, or heteroaromatic motif tends to appear in compounds with better permeability, lower clearance, or reduced safety risk. "
        "However, public databases mostly expose molecule-level records, assay-level measurements, or model-level predictions. This makes fragment-centric "
        "triage difficult, especially when users need to compare alternative motifs across multiple ADMET endpoints while retaining structural context."
    )
    add_paragraph(
        doc,
        "MedChem Fragment Atlas addresses this gap by building an attachment-aware dictionary of BRICS fragments and aggregating properties from parent "
        "molecules that contain each fragment. The atlas deliberately avoids claiming that a fragment has an intrinsic ADMET value. Instead, it estimates "
        "the distribution of parent-molecule measurements associated with a fragment, making the output useful for prioritizing hypotheses, selecting motifs "
        "for matched-series follow-up, and identifying fragments with sparse or uncertain evidence."
    )

    doc.add_heading("2. Results Overview", level=1)
    add_figure(
        doc,
        "figure_1_workflow.png",
        "Figure 1. MedChem Fragment Atlas workflow. ChEMBL molecules are validated, decomposed with RDKit BRICS, canonicalized while preserving dummy atoms and attachment labels, aggregated with parent-molecule descriptors and endpoints, and served through a DuckDB/FastAPI/React application.",
    )
    add_figure(
        doc,
        "figure_2_summary.png",
        f"Figure 2. Prototype atlas summary. The current build includes {c['molecules']:,} ChEMBL parent molecules, {c['fragments']:,} ChEMBL fragments, {c['molecule_fragments']:,} ChEMBL molecule-fragment mappings, and {c['tdc_admet_measurements']:,} clean TDC ADMET measurements.",
    )

    doc.add_heading("3. Methods", level=1)
    doc.add_heading("3.1 Molecule Preparation and Fragmentation", level=2)
    add_paragraph(
        doc,
        "The pipeline accepts either a ChEMBL SQLite database or a CSV file containing molecule identifiers and canonical SMILES. Molecules are parsed and "
        "sanitized with RDKit. Invalid molecules are excluded and logged. BRICS decomposition is performed with RDKit using leaf fragments only. Fragment "
        "SMILES are canonicalized with BRICS dummy atoms and attachment labels retained. A separate display_smiles representation normalizes dummy labels "
        "for visual display. Fragments with fewer than three heavy atoms are removed by default."
    )
    doc.add_heading("3.2 Parent-Molecule Descriptor Aggregation", level=2)
    add_paragraph(
        doc,
        "For parent molecules, missing descriptors are computed with RDKit: molecular weight, clogP, TPSA, hydrogen-bond donors and acceptors, rotatable "
        "bonds, ring count, aromatic ring count, fraction sp3, and QED. These descriptors are aggregated across parent molecules containing each fragment "
        "using mean, median, standard deviation, and percentile summaries. These values describe the parent molecules associated with a fragment, not the "
        "isolated fragment."
    )
    doc.add_heading("3.3 Clean TDC ADMET Mapping", level=2)
    add_paragraph(
        doc,
        "Clean ADMET benchmarks are imported from Therapeutics Data Commons ADME and Tox tasks. Endpoints are normalized into medicinal chemistry categories "
        "including Papp, absorption, clearance, half-life, plasma protein binding, volume of distribution, solubility, LogD, hERG safety, LD50, and DILI. "
        "Classification endpoints are retained as binary or safety-class values, while continuous endpoints preserve task-specific units. TDC molecules are "
        "independently decomposed with the same BRICS settings and joined to clean endpoint measurements."
    )
    doc.add_heading("3.4 Bootstrap Uncertainty", level=2)
    add_paragraph(
        doc,
        f"To quantify statistical uncertainty, we bootstrap fragment-endpoint medians for groups with at least 20 measurements. The current run produced "
        f"{c['bootstrap_ci']:,} fragment-endpoint confidence interval rows. Each bootstrap result reports the observed median, percentile 95% confidence "
        "interval, interval width, molecule count, measurement count, and contributing TDC tasks."
    )
    doc.add_heading("3.5 Matched-Context Validation", level=2)
    add_paragraph(
        doc,
        "To partially control for whole-molecule confounding, we perform a nearest-neighbor matched-context analysis. For each fragment-endpoint-task group, "
        "molecules containing the fragment are matched to molecules from the same TDC task that do not contain the fragment. Matching uses standardized MW, "
        "clogP, TPSA, HBD, HBA, and rotatable bond count. Endpoint differences are transformed so positive values indicate a more favorable association "
        "when the endpoint has a clear high- or low-preferred direction."
    )

    doc.add_heading("4. Fragment Atlas Results", level=1)
    add_figure(
        doc,
        "figure_3_top_fragments.png",
        "Figure 3. Most frequent BRICS fragments in the ChEMBL prototype. Fragment labels are display SMILES with normalized attachment sites.",
    )
    add_figure(
        doc,
        "figure_4_descriptor_heatmap.png",
        "Figure 4. Parent-molecule descriptor medians by fragment. Descriptor values are aggregated from parent molecules containing each fragment and should not be interpreted as isolated-fragment properties.",
    )
    add_figure(
        doc,
        "figure_5_admet_heatmap.png",
        "Figure 5. Clean TDC ADMET endpoint heatmap. Median endpoint values are mapped onto BRICS fragments and colored by medicinal chemistry desirability targets.",
    )

    doc.add_heading("Table 1. Top Fragment Examples", level=2)
    add_table(
        doc,
        ["Fragment", "Parents", "ADMET rows", "MW median", "LogP median", "TPSA median", "QED median"],
        [
            [
                row["display_smiles"],
                row["parent_count"],
                row["admet_measurement_count"],
                fmt(row["mw_median"]),
                fmt(row["clogp_median"]),
                fmt(row["tpsa_median"]),
                fmt(row["qed_median"]),
            ]
            for row in data["top_fragments"][:8]
        ],
    )

    doc.add_heading("5. Validation Experiments", level=1)
    add_figure(
        doc,
        "figure_8_bootstrap_ci.png",
        f"Figure 6. Bootstrap confidence intervals for selected fragment ADMET medians. The full output contains {c['bootstrap_ci']:,} fragment-endpoint rows.",
    )
    add_figure(
        doc,
        "figure_9_matched_context_validation.png",
        f"Figure 7. Matched-context validation. Fragment-containing parent molecules are compared with nearest-neighbor controls from the same TDC task. The full output contains {c['matched_context']:,} matched-context rows.",
    )

    doc.add_heading("Table 2. Representative Matched-Context Effects", level=2)
    add_table(
        doc,
        ["Fragment", "Endpoint", "Task", "Pairs", "Case median", "Control median", "Favorable delta"],
        [
            [
                row["display_smiles"],
                f"{row['standard_type']} ({row.get('standard_units') or 'unitless'})",
                row["tdc_task"],
                row["n_pairs"],
                fmt(row["case_median"]),
                fmt(row["control_median"]),
                fmt(row["favorable_median_delta"]),
            ]
            for row in data["matched_examples"][:8]
        ],
    )

    doc.add_heading("6. Case Studies", level=1)
    add_paragraph(
        doc,
        "The case studies combine fragment-level endpoint distributions, bootstrap uncertainty, and matched-context effect sizes. They are designed for "
        "medicinal chemistry triage rather than causal claims. Each association should be treated as a hypothesis to test in matched chemical series."
    )
    add_figure(
        doc,
        "figure_10_case_study_panels.png",
        "Figure 8. Publication-style case-study panels summarizing favorable and unfavorable fragment ADMET associations.",
    )
    add_table(
        doc,
        ["Fragment", "Endpoint", "Task", "Pairs", "Direction", "Delta", "Bootstrap CI"],
        [
            [
                row["display_smiles"],
                f"{row['standard_type']} ({row.get('standard_units') or 'unitless'})",
                row["tdc_task"],
                row["n_pairs"],
                row["case_direction"],
                fmt(row["favorable_median_delta"]),
                f"{fmt(row['ci_low'])} to {fmt(row['ci_high'])}",
            ]
            for row in data["case_studies"]
        ],
    )

    doc.add_heading("7. Web Application", level=1)
    add_paragraph(
        doc,
        "The companion web application exposes fragment search, endpoint-specific heatmaps, clickable median distributions, standard deviation summaries, "
        "representative parent molecules, and a two-fragment comparison page with side-by-side ADMET box plots. The interface repeatedly states that "
        "properties are aggregated from parent molecules containing each fragment."
    )
    add_figure(
        doc,
        "figure_7_web_interface.png",
        "Figure 9. Web interface for searching fragments and inspecting parent-molecule endpoint summaries.",
    )

    doc.add_heading("8. Discussion", level=1)
    add_paragraph(
        doc,
        "The atlas provides a practical way to navigate fragment-associated ADMET evidence. Bootstrap intervals expose uncertainty for low-count fragments, "
        "while matched-context validation reduces, but does not eliminate, confounding from whole-molecule properties. The results highlight that some "
        "associations remain strongly context-dependent: a fragment can appear favorable in one endpoint, unfavorable in another, or ambiguous when measured "
        "under a different task definition."
    )
    add_paragraph(
        doc,
        "The strongest use case is prioritization. A medicinal chemist can compare candidate motifs, identify endpoints where one motif has a broader or "
        "shifted distribution, and then inspect representative parent molecules. This can support ideation before more rigorous modeling, matched molecular "
        "pair analysis, or prospective synthesis."
    )

    doc.add_heading("9. Limitations", level=1)
    add_bullets(
        doc,
        [
            "Fragment-associated ADMET summaries are observational and should not be interpreted as causal fragment effects.",
            "BRICS fragmentation captures synthetically meaningful cuts but does not represent every chemically relevant motif or matched replacement.",
            "TDC tasks differ in endpoint definition, units, and assay context; endpoints are therefore stratified by task and units where possible.",
            "Nearest-neighbor matching controls simple physicochemical descriptors but not all scaffold, target, series, assay, or publication biases.",
            "Binary safety and absorption endpoints can produce degenerate medians and narrow bootstrap intervals even when uncertainty remains practically important.",
        ],
    )

    doc.add_heading("10. Conclusion", level=1)
    add_paragraph(
        doc,
        "MedChem Fragment Atlas is an open, reproducible, and interactive fragment-centric view of parent-molecule ADMET associations. By combining RDKit "
        "BRICS decomposition, clean TDC endpoint mapping, bootstrap uncertainty, matched-context validation, and web-based box-plot comparison, the atlas "
        "turns public molecule-level data into a practical hypothesis generator for medicinal chemistry optimization."
    )

    doc.add_heading("Code and Data Availability", level=1)
    add_paragraph(
        doc,
        "The project contains the full FastAPI backend, React frontend, DuckDB schema, and reproducible pipeline scripts. Source data are derived from "
        "ChEMBL and Therapeutics Data Commons. Generated artifacts include Parquet tables, DuckDB serving tables, SVG/PNG figures, case-study tables, and "
        "this manuscript draft."
    )

    doc.add_heading("References", level=1)
    refs = [
        "Bento et al. The ChEMBL bioactivity database: an update. Nucleic Acids Research, 2014.",
        "Mendez et al. ChEMBL: towards direct deposition of bioassay data. Nucleic Acids Research, 2019.",
        "Landrum G. RDKit: Open-source cheminformatics software.",
        "Degen et al. On the art of compiling and using 'drug-like' chemical fragment spaces. ChemMedChem, 2008.",
        "Huang et al. Therapeutics Data Commons: machine learning datasets and tasks for drug discovery and development. NeurIPS Datasets and Benchmarks, 2021.",
        "Bickerton et al. Quantifying the chemical beauty of drugs. Nature Chemistry, 2012.",
        "Lipinski et al. Experimental and computational approaches to estimate solubility and permeability in drug discovery and development settings. Advanced Drug Delivery Reviews, 2001.",
        "Veber et al. Molecular properties that influence the oral bioavailability of drug candidates. Journal of Medicinal Chemistry, 2002.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
